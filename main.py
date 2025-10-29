import os
import io
import re
import logging
from io import BytesIO
from typing import Optional, Tuple
from datetime import datetime, timedelta
import secrets

import requests
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI
from PIL import Image

# =========================
# Logs
# =========================
logging.basicConfig(level=logging.INFO, format="%(asctime)s — %(levelname)s — %(message)s")
logger = logging.getLogger(__name__)

REPORT_COLOR   = RGBColor(133, 78, 197)
HEADING_COLOR  = RGBColor(85, 54, 185)

# =========================
# Assets
# =========================
ASSETS_DIR = "assets"
DEFAULT_PORTADA_PATH       = os.path.join(ASSETS_DIR, "portada.png")
DEFAULT_CONTRAPORTADA_PATH = os.path.join(ASSETS_DIR, "contraportada.png")
DEFAULT_LOGO_PATH          = os.path.join(ASSETS_DIR, "logo.png")

app = FastAPI(title="Formateador de informes Dipli")

# =========================
# CORS
# =========================
ALLOWED_ORIGINS = [
    "https://www.dipli.ai",
    "https://dipli.ai",
    "https://isagarcivill09.wixsite.com/turop",
    "https://isagarcivill09.wixsite.com/turop/tienda",
    "https://isagarcivill09-wixsite-com.filesusr.com",
    "https://www.dipli.ai/preparaci%C3%B3n",
    "https://www-dipli-ai.filesusr.com",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    allow_origin_regex=r"https://.*\.filesusr\.com",
)

# =========================
# Descargas temporales
# =========================
DOWNLOADS: dict[str, tuple[bytes, str, str, datetime]] = {}
DOWNLOAD_TTL_SECS = 900  # 15 min
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

def cleanup_downloads() -> None:
    now = datetime.utcnow()
    expired = [t for t, (_, _, _, exp) in DOWNLOADS.items() if exp <= now]
    for t in expired:
        DOWNLOADS.pop(t, None)

def register_download(data: bytes, filename: str, media_type: str) -> str:
    cleanup_downloads()
    token = secrets.token_urlsafe(16)
    expires_at = datetime.utcnow() + timedelta(seconds=DOWNLOAD_TTL_SECS)
    DOWNLOADS[token] = (data, filename, media_type, expires_at)
    return token

def ensure_default_assets() -> None:
    try:
        os.makedirs(ASSETS_DIR, exist_ok=True)
        if not os.path.exists(DEFAULT_PORTADA_PATH):
            Image.new("RGB", (1200, 1600), (133, 78, 197)).save(DEFAULT_PORTADA_PATH, format="PNG")
        if not os.path.exists(DEFAULT_CONTRAPORTADA_PATH):
            Image.new("RGB", (1200, 1600), (85, 54, 185)).save(DEFAULT_CONTRAPORTADA_PATH, format="PNG")
        if not os.path.exists(DEFAULT_LOGO_PATH):
            img = Image.new("RGB", (600, 600), (255, 255, 255))
            try:
                from PIL import ImageDraw
                draw = ImageDraw.Draw(img)
                draw.ellipse((100, 100, 500, 500), fill=(133, 78, 197))
            except Exception:
                pass
            img.save(DEFAULT_LOGO_PATH, format="PNG")
    except Exception as e:
        logger.warning(f"No se pudieron preparar assets por defecto: {e}")

ensure_default_assets()

# =========================
# OpenAI helpers
# =========================
USE_TIKTOKEN = False
MODEL_MAX_TOKENS = 8192
try:
    import tiktoken
    ENCODING = tiktoken.encoding_for_model("gpt-4")
    USE_TIKTOKEN = True
except Exception:
    ENCODING = None

def trim_to_fit(text: str, reserved_output: int = 700) -> str:
    if USE_TIKTOKEN and ENCODING:
        tokens = ENCODING.encode(text)
        max_input = max(MODEL_MAX_TOKENS - reserved_output - 100, 0)
        return ENCODING.decode(tokens[: max_input])
    approx_chars_per_token = 4
    max_input_tokens = max(MODEL_MAX_TOKENS - reserved_output - 100, 0)
    return text[: max_input_tokens * approx_chars_per_token]

def call_gpt(api_key: str, prompt: str, user_input: str, max_tokens: int = 700) -> str:
    """Devuelve texto del modelo o "" si no hay API key o hay error."""
    if not api_key or not api_key.strip():
        logger.info("OpenAI API key ausente -> se omite generación.")
        return ""
    try:
        client = OpenAI(api_key=api_key.strip())
        trimmed_input = trim_to_fit(user_input, reserved_output=max_tokens)
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": trimmed_input},
            ],
            temperature=0.5,
            max_tokens=max_tokens,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        logger.error(f"Error llamando a OpenAI: {e}")
        return ""

# =========================
# Prompts
# =========================
PROMPT_RESUMEN = """
Actúa como un experto en redacción ejecutiva y análisis de informes cualitativos. Tu tarea es redactar un resumen ejecutivo profesional y conciso, basándote en el siguiente contenido del informe.

**Requisitos:**
* Longitud: proporcional al documento original (~3% del total de palabras).
* Estructura: dos párrafos integrados.
* Contenido: objetivo, alcance, hallazgos, impacto y recomendaciones.
* Tono: formal, claro y accesible.
* Evitar: repeticiones, tecnicismos y explicaciones extensas.

"""
PROMPT_HALLAZGOS = """
Actúa como un experto en redacción ejecutiva y análisis cualitativo. A partir del siguiente documento, redacta una sección titulada "Principales Hallazgos" en español.

**Requisitos:**
* Enfócate en los hallazgos clave y sus implicaciones.
* Estructura: numeración y contexto breve por punto.
* Estilo: profesional, claro, útil para tomadores de decisiones.
* Evitar: recomendaciones, juicios, generalidades o conclusiones.
"""
# =========================
# Márgenes del cuerpo y helper
# =========================
BODY_TOP = Cm(2.5)
BODY_BOTTOM = Cm(2.5)
BODY_LEFT = Cm(2.5)
BODY_RIGHT = Cm(2.5)

def set_body_margins(section):
    section.top_margin = BODY_TOP
    section.bottom_margin = BODY_BOTTOM
    section.left_margin = BODY_LEFT
    section.right_margin = BODY_RIGHT

# =========================
# Helpers DOCX (estilos, portada, etc.)
# =========================
def modify_style(doc: Document, style_name: str, size_pt: int,
                 bold: bool = False, italic: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    font = doc.styles[style_name].font
    font.name = "Century Gothic"
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color
    pf = doc.styles[style_name].paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

def insert_cover_page(doc_out: Document, portada_path) -> None:
    # Portada en sección 0 a sangrado 0
    sec0 = doc_out.sections[0]
    sec0.top_margin = sec0.bottom_margin = Cm(0)
    sec0.left_margin = sec0.right_margin = Cm(0)

    pw, ph = sec0.page_width, sec0.page_height
    try:
        if isinstance(portada_path, list):
            resp = requests.get(portada_path[0], timeout=15)
            img = BytesIO(resp.content)
        else:
            img = portada_path
    except Exception as e:
        logger.error(f"Error portada: {e}")
        return

    p = doc_out.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run().add_picture(img, width=pw, height=ph)

    # NUEVA sección de cuerpo con márgenes normales
    body_sec = doc_out.add_section(WD_SECTION.NEW_PAGE)
    set_body_margins(body_sec)

def insert_contraportada_body(doc_out: Document, contraportada_path) -> None:
    # Sección de contraportada a sangrado 0
    cover_sec = doc_out.add_section(WD_SECTION.NEW_PAGE)
    cover_sec.top_margin = cover_sec.bottom_margin = Cm(0)
    cover_sec.left_margin = cover_sec.right_margin = Cm(0)

    pw, ph = cover_sec.page_width, cover_sec.page_height
    try:
        if isinstance(contraportada_path, str):
            resp = requests.get(contraportada_path, timeout=15)
            img = BytesIO(resp.content)
        else:
            img = contraportada_path
    except Exception as e:
        logger.error(f"Error contraportada: {e}")
        return

    p = doc_out.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run().add_picture(img, width=pw, height=ph)

    # NUEVA sección de cuerpo con márgenes normales
    body_sec = doc_out.add_section(WD_SECTION.NEW_PAGE)
    set_body_margins(body_sec)

def insert_footer_logo(doc: Document, logo_source) -> None:
    sec = doc.sections[-1]
    footer = sec.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    try:
        if isinstance(logo_source, str):
            resp = requests.get(logo_source, timeout=15)
            img = BytesIO(resp.content)
        else:
            img = logo_source
        p.add_run().add_picture(img, width=Inches(1.5))
    except Exception as e:
        logger.error(f"Error footer logo: {e}")

def add_table_of_contents(paragraph) -> None:
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-5" \\h \\z \\u')
    paragraph._p.append(fld)


# ====== verbatims centrados (_"..."_ y *"..."*) ======
def format_text_block(doc: Document, texto: str, color=RGBColor(133, 78, 197)) -> None:
    def _last_is_blank() -> bool:
        return bool(doc.paragraphs) and doc.paragraphs[-1].text.strip() == ""
    def _add_blank_once():
        if not _last_is_blank():
            bp = doc.add_paragraph("")
            pf = bp.paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(0)

    tokens = re.split(r'(\*\*[^*]+\*\*|_"[^"]+"_|\*"[^"]+"\*|_[^_]+_|[*][^*]+[*])', texto)
    p = None
    def ensure_p():
        nonlocal p
        if p is None:
            p = doc.add_paragraph("", style="Normal")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            pf = p.paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(0)
        return p

    for token in tokens:
        if not token:
            continue
        # _"..."_
        if token.startswith('_"') and token.endswith('"_'):
            _add_blank_once()
            vp = doc.add_paragraph("", style="Normal")
            vp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            vpf = vp.paragraph_format
            vpf.space_before = Pt(0); vpf.space_after = Pt(0)
            run = vp.add_run(token[2:-2]); run.bold = True; run.italic = True; run.font.color.rgb = color
            _add_blank_once(); p = None; continue
        # *"..."*
        if token.startswith('*"') and token.endswith('"*'):
            _add_blank_once()
            vp = doc.add_paragraph("", style="Normal")
            vp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            vpf = vp.paragraph_format
            vpf.space_before = Pt(0); vpf.space_after = Pt(0)
            run = vp.add_run(token[2:-2]); run.bold = True; run.italic = True; run.font.color.rgb = color
            _add_blank_once(); p = None; continue
        # **negrita**
        if token.startswith("**") and token.endswith("**"):
            r = ensure_p().add_run(token[2:-2]); r.bold = True; continue
        # _itálica_
        if token.startswith("_") and token.endswith("_"):
            r = ensure_p().add_run(token[1:-1]); r.italic = True; continue
        # *negrita+itálica* (coloreada)
        if token.startswith("*") and token.endswith("*"):
            r = ensure_p().add_run(token[1:-1]); r.bold = True; r.italic = True; r.font.color.rgb = color; continue
        # texto plano
        ensure_p().add_run(token)

# ====== títulos duplicados ======
def _normalize_title(s: str) -> str:
    s = re.sub(r"^[#\s]+", "", s.replace("\xa0", " "))
    s = s.strip("*_ :\t-—").lower()
    return s
def _is_duplicate_section_title(texto: str) -> bool:
    norm = _normalize_title(texto)
    return norm.startswith("resumen ejecutivo") or norm.startswith("principales hallazgos")

# ====== detección robusta de headings estilo Markdown ======
def parse_hash_heading(raw: str) -> Tuple[Optional[int], Optional[str]]:
    """
    Devuelve (level, title) si el párrafo es '### Título', '####   Título', etc.
    Tolera NBSP (\xa0) y espacios múltiples. Solo reconoce 3 o más '#'.
    """
    if not raw:
        return None, None
    s = raw.replace("\xa0", " ").lstrip()
    m = re.match(r'^(#{3,})\s+(.*\S)\s*$', s)
    if not m:
        return None, None
    level = len(m.group(1))
    title = m.group(2).strip()
    return level, title

# =========================
# Generación de informe
# =========================
def generate_report(api_key: str,
                    input_doc_bytes: bytes,
                    portada_bytes: Optional[bytes],
                    contraportada_bytes: Optional[bytes],
                    logo_bytes: Optional[bytes],
                    use_defaults: bool,
                    filename_hint: str) -> bytes:
    doc = Document(BytesIO(input_doc_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text  = "\n\n".join(paragraphs)

    # Selección de assets con fallback por-asset
    if use_defaults:
        with open(DEFAULT_PORTADA_PATH, "rb") as f: portada_path = BytesIO(f.read())
        with open(DEFAULT_CONTRAPORTADA_PATH, "rb") as f: contraportada_path = BytesIO(f.read())
        with open(DEFAULT_LOGO_PATH, "rb") as f: logo_path = BytesIO(f.read())
    else:
        if portada_bytes:
            portada_path = BytesIO(portada_bytes)
        else:
            with open(DEFAULT_PORTADA_PATH, "rb") as f: portada_path = BytesIO(f.read())
        if contraportada_bytes:
            contraportada_path = BytesIO(contraportada_bytes)
        else:
            with open(DEFAULT_CONTRAPORTADA_PATH, "rb") as f: contraportada_path = BytesIO(f.read())
        if logo_bytes:
            logo_path = BytesIO(logo_bytes)
        else:
            with open(DEFAULT_LOGO_PATH, "rb") as f: logo_path = BytesIO(f.read())

    doc_out = Document()
    insert_cover_page(doc_out, portada_path)         # crea sección de cuerpo con márgenes
    set_body_margins(doc_out.sections[-1])           # cinturón y tirantes

    modify_style(doc_out, 'Normal',    12)
    modify_style(doc_out, 'Heading 1', 20, bold=True,  color=HEADING_COLOR)
    modify_style(doc_out, 'Heading 2', 16, bold=True,  color=HEADING_COLOR)
    modify_style(doc_out, 'Heading 3', 14, bold=True,  color=HEADING_COLOR)
    modify_style(doc_out, 'Heading 4', 12, bold=True,  color=HEADING_COLOR)
    modify_style(doc_out, 'Heading 5', 20, bold=True,  color=REPORT_COLOR)

    doc_out.add_paragraph("Tabla de contenidos", style="Heading 5")
    add_table_of_contents(doc_out.add_paragraph())
                        
    # Resumen (opcional)
    resumen = call_gpt(api_key, PROMPT_RESUMEN, full_text[:10000], 500)
    if resumen:
        doc_out.add_paragraph("Resumen Ejecutivo", style="Heading 1")
        doc_out.add_paragraph(resumen, style="Normal")

    # Hallazgos (opcional)
    hallazgos = call_gpt(api_key, PROMPT_HALLAZGOS, full_text[:10000], 700)
    if hallazgos:
        doc_out.add_paragraph("Principales Hallazgos", style="Heading 1")
        for line in hallazgos.split("\n"):
            t = line.strip()
            if not t or _is_duplicate_section_title(t):
                continue
            format_text_block(doc_out, t, color=REPORT_COLOR)

    # Cuerpo original (con detección robusta de headings)
    for para in paragraphs:
        t = para.strip()
        if not t or _is_duplicate_section_title(t):
            continue

        level, title = parse_hash_heading(t)
        if level:
            if level == 3:
                # Capítulo principal: contraportada + Heading 1
                insert_contraportada_body(doc_out, contraportada_path)  # deja nueva sección con márgenes
                doc_out.add_paragraph(title, style="Heading 1")
            elif level == 4:
                doc_out.add_paragraph(title, style="Heading 2")
            elif level == 5:
                doc_out.add_paragraph(title, style="Heading 3")
            else:
                doc_out.add_paragraph(title, style="Heading 4")
            continue  # no imprimir "### ..."

        # Contenido normal
        format_text_block(doc_out, t, color=REPORT_COLOR)

    insert_footer_logo(doc_out, logo_path)

    out_bytes = io.BytesIO()
    doc_out.save(out_bytes)
    out_bytes.seek(0)
    return out_bytes.getvalue()

# ---------- Endpoints ----------
@app.post("/generate-report-link")
async def generate_report_link(
    request: Request,
    file: UploadFile = File(..., description="Archivo .docx base"),
    openai_api_key: str = Form("", description="Tu OpenAI API Key (opcional)"),
    usar_personalizadas: str = Form("false"),
    portada: UploadFile | None = File(None),
    contraportada: UploadFile | None = File(None),
    logo: UploadFile | None = File(None),
):
    try:
        # normalizar flag a bool (acepta "true", "1", "yes", "on")
        usar_personalizadas_bool = str(usar_personalizadas).strip().lower() in ("true", "1", "yes", "on")

        if not file.filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="Debes subir un archivo .docx válido.")

        base_bytes = await file.read()

        # Leer bytes si vienen (y si el switch está activo)
        portada_bytes = await portada.read() if (usar_personalizadas_bool and portada is not None and hasattr(portada, "read")) else None
        contraportada_bytes = await contraportada.read() if (usar_personalizadas_bool and contraportada is not None and hasattr(contraportada, "read")) else None
        logo_bytes = await logo.read() if (usar_personalizadas_bool and logo is not None and hasattr(logo, "read")) else None

        result = generate_report(
            api_key=openai_api_key,
            input_doc_bytes=base_bytes,
            portada_bytes=portada_bytes,              # puede ser None
            contraportada_bytes=contraportada_bytes,  # puede ser None
            logo_bytes=logo_bytes,                    # puede ser None
            use_defaults=not usar_personalizadas_bool,  # si el switch está apagado, forzar defaults
            filename_hint=file.filename,
        )

        final_name = file.filename.replace(".docx", "") + "_INFORME_FINAL.docx"
        token = register_download(result, final_name, DOCX_MEDIA_TYPE)

        base_url = str(request.base_url).rstrip("/")
        return {"download_url": f"{base_url}/download/{token}", "expires_in_seconds": DOWNLOAD_TTL_SECS}

    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Fallo en /generate-report-link")
        raise HTTPException(status_code=500, detail=f"Error al generar el informe: {e}")

@app.get("/download/{token}")
def download_token(token: str):
    cleanup_downloads()
    item = DOWNLOADS.get(token)
    if not item:
        raise HTTPException(status_code=404, detail="Link expirado o inválido")
    data, filename, media_type, exp = item
    if exp <= datetime.utcnow():
        DOWNLOADS.pop(token, None)
        raise HTTPException(status_code=410, detail="Link expirado")
    headers = {"Content-Disposition": f'attachment; filename="{filename}"', "Cache-Control": "no-store"}
    return StreamingResponse(io.BytesIO(data), media_type=media_type, headers=headers)

@app.get("/")
def root():
    return {"status": "ok", "service": "dipli-docx-generator"}

@app.get("/health")
def health():
    return {"status": "healthy", "message": "API funcionando correctamente"}

